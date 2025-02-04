#version 3
import json
import re
import uuid
import oci
from flask import Flask, request, jsonify, render_template, send_from_directory
from flask_cors import CORS
from flask import session
from werkzeug.utils import secure_filename
from opensearchpy import OpenSearch, RequestsHttpConnection
from requests.auth import HTTPBasicAuth
import os
import fitz  
import io
import redis
import docx
from PyPDF2 import PdfReader
import traceback
from config import Config
import numpy as np
from redis import Redis
import tiktoken
import docx
from docx import Document
from flask_session import Session
from operator import itemgetter
import concurrent.futures




app = Flask(__name__, static_folder='static')
CORS(app)
app.config.from_object(Config)
MAX_WORKERS = 10 

app.config['SESSION_TYPE'] = 'filesystem'
app.config['SESSION_PERMANENT'] = False
app.config['SESSION_USE_SIGNER'] = True
app.config['SECRET_KEY'] = 'your_secret_key_here'  
Session(app)

default_config = oci.config.from_file(app.config['OCI_CONFIG_FILE'], app.config['OCI_DEFAULT_PROFILE'])
chicago_config = oci.config.from_file(app.config['OCI_CONFIG_FILE'], app.config['OCI_CHICAGO_PROFILE'])
object_storage_client = oci.object_storage.ObjectStorageClient(chicago_config)

endpoint =  app.config['OCI_ENDPOINT']

opensearch_client = OpenSearch(
    hosts=[{'host': app.config['OPENSEARCH_HOST'], 'port': app.config['OPENSEARCH_PORT']}],
    http_auth=(app.config['OPENSEARCH_USER'], app.config['OPENSEARCH_PASSWORD']),
    use_ssl=True,
    verify_certs=False,
    connection_class=RequestsHttpConnection
)


redis_client = redis.StrictRedis(
    host=app.config['REDIS_HOST'],  
    port=app.config['REDIS_PORT'],  
    decode_responses=True  
)




generative_ai_client = oci.generative_ai.GenerativeAiClient(chicago_config)
generative_ai_inference_client = oci.generative_ai_inference.GenerativeAiInferenceClient(
    config=chicago_config, 
    service_endpoint=endpoint,
    retry_strategy=oci.retry.NoneRetryStrategy(),
    timeout=(10, 240)
)




def create_opensearch_index(index_name):
    index_body = {
        "settings": {
            "index": {
                "knn": True,
                "knn.algo_param.ef_search": 100,
                "number_of_shards": 1,
                "number_of_replicas": 0
            }
        },
        "mappings": {
            "properties": {
                "text": {
                    "type": "text"
                },
                "passage_embedding": {
                    "type": "knn_vector",
                    "dimension": 384, 
                    "method": {
                        "name": "hnsw",
                        "space_type": "l2",
                        "engine": "nmslib",
                        "parameters": {
                            "ef_construction": 128,
                            "m": 24
                        }
                    }
                }
            }
        }
    }
    opensearch_client.indices.create(index=index_name, body=index_body)





def chunk_text(text, max_tokens=500): 
    encoding = tiktoken.get_encoding("cl100k_base")
    tokens = encoding.encode(text)
    
    chunks = []
    current_chunk = []
    current_length = 0

    for token in tokens:
        if current_length + 1 < max_tokens:
            current_chunk.append(token)
            current_length += 1
        else:
            chunk_text = encoding.decode(current_chunk)
            chunks.append(chunk_text)
            print(f"Chunk created with {current_length} tokens")
            current_chunk = [token]
            current_length = 1

    if current_chunk:
        chunk_text = encoding.decode(current_chunk)
        chunks.append(chunk_text)
        print(f"Final chunk created with {current_length} tokens")

    return chunks

def call_embedding_api(text):
    try:
        chunks = chunk_text(text, max_tokens=500)  # Use 500 
        embeddings = []

        for i, chunk in enumerate(chunks):
            print(f"Processing chunk {i+1}/{len(chunks)}")
            encoding = tiktoken.get_encoding("cl100k_base")
            token_count = len(encoding.encode(chunk))
            print(f"Chunk token count: {token_count}")
            
            if token_count > 512:
                print(f"Warning: Chunk {i+1} has {token_count} tokens, which exceeds the limit.")
                chunk = encoding.decode(encoding.encode(chunk)[:512])
                print(f"Truncated chunk to {len(encoding.encode(chunk))} tokens")

            embed_text_details = oci.generative_ai_inference.models.EmbedTextDetails(
                serving_mode=oci.generative_ai_inference.models.OnDemandServingMode(model_id="cohere.embed-multilingual-light-v3.0"),
                inputs=[chunk],
                truncate="NONE",
                compartment_id=app.config['OCI_COMPARTMENT_ID']
            )
            response = generative_ai_inference_client.embed_text(embed_text_details)
            embeddings.append(response.data.embeddings[0])

        if len(embeddings) > 1:
            return aggregate_embeddings(embeddings)
        return embeddings[0]
    except oci.exceptions.ServiceError as e:
        print(f"Error in embedding API call: {str(e)}")
        return None
    except Exception as e:
        print(f"Unexpected error in call_embedding_api: {str(e)}")
        return None


def generate_embeddings_for_chunks(chunks):
    embeddings = []
    for chunk in chunks:
        embedding = call_embedding_api(chunk)
        if embedding is not None:
            embeddings.append(embedding)
    return embeddings

def aggregate_embeddings(embeddings):
    if not embeddings:
        return None
    aggregated_embedding = np.mean(embeddings, axis=0)
    return aggregated_embedding.tolist()



@app.route('/')
def index():
    recommended_params = {
        "temperature": 0.7,
        "max_tokens": 500,
        "frequency_penalty": 0.1,
        "top_p": 0.9,
        "top_k": 40
    }
    return render_template('index.html', recommended_params=recommended_params)

@app.route('/static/<path:path>')
def send_static(path):
    return send_from_directory('static', path)

@app.route('/list-models', methods=['GET'])
def list_models():
    try:
        list_models_response = generative_ai_client.list_models(
            compartment_id=app.config['OCI_COMPARTMENT_ID'],
            limit=100,
            sort_order="ASC",
            sort_by="displayName")
        
        models = list_models_response.data.items
        
        chat_models = [model for model in models if 'CHAT' in model.capabilities]
        
        print("Available CHAT models:")
        for model in chat_models:
            print(f"ID: {model.id}")
            print(f"Display Name: {model.display_name}")
            print(f"Capabilities: {model.capabilities}")
            
        
        return jsonify([{
            "id": model.id,
            "displayName": model.display_name,
            "capabilities": model.capabilities
        } for model in chat_models])
    except Exception as e:
        error_message = f"Error fetching models: {str(e)}\n{traceback.format_exc()}"
        app.logger.error(error_message)
        print(error_message)  
        return jsonify({"error": error_message}), 500

def clean_extracted_text(text):
    text = re.sub(r'\s+', ' ', text).strip()  
    text = re.sub(r'[^\x00-\x7F]+', ' ', text)  
    return text

def extract_text_from_pdf(pdf_bytes):
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        text = ""
        for page in doc:
            text += page.get_text()
        return clean_extracted_text(text)
    except Exception as e:
        print(f"Error extracting text from PDF: {str(e)}")
        return ""

def extract_text_from_docx(docx_bytes):
    try:
       
        docx_file = io.BytesIO(docx_bytes)
        doc = Document(docx_file)
        
        
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        
        
        return '\n'.join(full_text)
    except Exception as e:
        print(f"Error extracting text from DOCX: {str(e)}")
        return "Error: Unable to extract text from the DOCX file."

def process_file(file_path, file_content, filename):
    file_extension = os.path.splitext(filename)[1].lower()
    
    if file_extension == '.docx':
        content = extract_text_from_docx(io.BytesIO(file_content))
    elif file_extension == '.pdf':
        content = extract_text_from_pdf(io.BytesIO(file_content))
    else:
        return None  
    
    chunks = chunk_text(content)
    chunk_embeddings = generate_embeddings_for_chunks(chunks)
    aggregated_embedding = aggregate_embeddings(chunk_embeddings)
    
    if aggregated_embedding is None:
        return None
    
    doc_id = filename
    document_body = {"content": content, "passage_embedding": aggregated_embedding}
    opensearch_client.index(index="documents", id=doc_id, body=document_body)
    
    return filename



def upload_to_oci_bucket(file, object_name):
 
    try:
        print(f"Received file: {file.filename}, Content-Type: {file.content_type}")

        file_content = file.read()
        if not file_content:
            print("Error: File content is empty")
            return {"status": "error", "message": "File content is empty"}

        
        response = object_storage_client.put_object(
            namespace_name=app.config['OCI_NAMESPACE'],
            bucket_name=app.config['OCI_BUCKET_NAME'],
            object_name=object_name,
            put_object_body=file_content
        )

        if response.status == 200:
            print(f"Upload successful: {response.status} | Object: {object_name}")

            
            uploaded_files = list_objects_in_bucket()
            if object_name in uploaded_files:
                print(f"File {object_name} is confirmed in the bucket.")
                return {"status": "success", "object_name": object_name, "file_content": file_content}
            else:
                print(f"Error: File {object_name} is NOT in the bucket after upload!")
                return {"status": "error", "message": "File upload verification failed"}

        else:
            print(f"Upload failed with status {response.status}")
            return {"status": "error", "message": f"Upload failed with status {response.status}"}

    except Exception as e:
        print(f"Upload failed: {str(e)}")
        return {"status": "error", "message": str(e)}

def list_objects_in_bucket():
    """
    Lists all objects currently in the OCI bucket.
    """
    try:
        objects = object_storage_client.list_objects(
            namespace_name=app.config['OCI_NAMESPACE'],
            bucket_name=app.config['OCI_BUCKET_NAME']
        )

        uploaded_files = [obj.name for obj in objects.data.objects]
        print(f"Files currently in the bucket: {uploaded_files}")
        return uploaded_files
    except Exception as e:
        print(f"Error retrieving bucket contents: {str(e)}")
        return []

@app.route('/upload-to-oci', methods=['POST'])
def upload_to_oci():
    if 'files[]' not in request.files:
        return jsonify({"error": "No file part"}), 400

    files = request.files.getlist('files[]')
    if not files or all(file.filename.strip() == '' for file in files):
        return jsonify({"error": "No selected file"}), 400

    processed_files = []
    unsupported_files = []

    try:
        for file in files:
            filename = secure_filename(file.filename)

            if not filename.lower().endswith(('.pdf', '.docx')):
                unsupported_files.append(filename)
                continue  

            object_name = f"uploads/{filename}"
            upload_result = upload_to_oci_bucket(file, object_name)

            if upload_result["status"] != "success":
                app.logger.error(f"Failed to upload {filename} to OCI: {upload_result['message']}")
                continue

            app.logger.info(f"Uploaded {filename} to OCI bucket")

            file_content = upload_result["file_content"]

            if filename.lower().endswith('.pdf'):
                content = extract_text_from_pdf(io.BytesIO(file_content))
            elif filename.lower().endswith('.docx'):
                content = extract_text_from_docx(file_content)

            if not content.strip():
                app.logger.warning(f"No text extracted from {filename}, skipping indexing")
                continue

            print(f"Extracted content from {filename}: {content[:500]}")

            chunks = chunk_text(content)
            chunk_embeddings = generate_embeddings_for_chunks(chunks)
            aggregated_embedding = aggregate_embeddings(chunk_embeddings)

            try:
                doc_id = filename
                document_body = {
                    "content": content,
                    "passage_embedding": aggregated_embedding
                }
                response = opensearch_client.index(index="documents", id=doc_id, body=document_body)

                if response.get('result') not in ['created', 'updated']: 
                    app.logger.error(f"OpenSearch indexing issue for {filename} - Response: {response}")
                    continue

                app.logger.info(f"Indexed {filename} in OpenSearch")

            except Exception as os_error:
                app.logger.error(f"Failed to index {filename} in OpenSearch: {os_error}")
                continue

            processed_files.append(filename)

        return jsonify({
            "status": f"Processed {len(processed_files)} files successfully",
            "processed_files": processed_files,
            "unsupported_files": unsupported_files
        })

    except Exception as e:
        error_message = f"Error processing files: {str(e)}\n{traceback.format_exc()}"
        app.logger.error(error_message)
        return jsonify({"error": error_message}), 500


def clean_text(text):
    """
    Preprocess the text by removing unnecessary symbols, formatting artifacts, and noise.
    """
    text = re.sub(r'\s+', ' ', text)  
    text = re.sub(r'[^\w\s.,!?]', '', text)  
    return text.strip()


def estimate_tokens(text):
    """Estimate the number of tokens in a given text."""
    encoding = tiktoken.get_encoding("cl100k_base")
    return len(encoding.encode(text))

def select_relevant_content(contents, query, max_tokens):
    query_tokens = estimate_tokens(query)
    available_tokens = max_tokens - query_tokens - 200  

    selected_contents = []
    total_tokens = 0

    
    for content in contents:
        content_tokens = estimate_tokens(content)
        if total_tokens + content_tokens <= available_tokens:
            selected_contents.append(content)
            total_tokens += content_tokens
        else:
            break

    return "\n\n".join(selected_contents)


def get_chat_history(conversation_id, max_messages=5):
    try:
        chat_history = redis_client.lrange(f"chat:{conversation_id}", -max_messages, -1)
        return [tuple(json.loads(message)) for message in chat_history]
    except Exception as e:
        print(f"Error retrieving chat history from Redis: {str(e)}")
        return []

def store_chat_message(conversation_id, role, content):
    try:
        redis_client.rpush(f"chat:{conversation_id}", json.dumps((role, content)))
        redis_client.expire(f"chat:{conversation_id}", 3600)  # Set expiration to 1 hour
        print(f"Successfully stored message for conversation {conversation_id}")
    except Exception as e:
        print(f"Error storing message in Redis: {str(e)}")

def generate_conversation_id():
    return str(uuid.uuid4())

def count_tokens(text):
    encoding = tiktoken.get_encoding("cl100k_base")
    return len(encoding.encode(text))

def truncate_text(text, max_tokens):
    encoding = tiktoken.get_encoding("cl100k_base")
    tokens = encoding.encode(text)
    if len(tokens) > max_tokens:
        return encoding.decode(tokens[:max_tokens])
    return text

@app.route('/query', methods=['POST'])
def query():
    try:
        data = request.json
        query_text = data.get('query')
        conversation_id = data.get('conversation_id')
        model_id = data.get('model_id')
        model_display_name = data.get('model_display_name', '').lower()
        max_tokens = data.get('max_tokens', 600)
        temperature = data.get('temperature', 1.0)
        frequency_penalty = data.get('frequency_penalty', 0.0)
        top_p = data.get('top_p', 0.75)
        top_k = data.get('top_k', -1)

        if not query_text or not model_id:
            return jsonify({"error": "Missing required parameters: 'query' and 'model_id' are required"}), 400

        if not conversation_id:
            conversation_id = generate_conversation_id()
            print(f"Generated new conversation ID: {conversation_id}")
        else:
            print(f"Using existing conversation ID: {conversation_id}")

        print(f"Processing query for conversation ID: {conversation_id}")
        print(f"User query: {query_text}")

       
        chat_history = get_chat_history(conversation_id)
        print(f"Retrieved chat history: {chat_history}")
        chat_history_text = "\n".join([f"{role.upper()}: {content}" for role, content in chat_history])
        chat_history_text += f"\nUSER: {query_text}"

        
        query_embedding = call_embedding_api(query_text)
        if query_embedding is None:
            return jsonify({"error": "Failed to generate query embedding"}), 500

        vector = query_embedding if isinstance(query_embedding, list) else [float(query_embedding)]

        # Retrieve all relevant documents using a larger 'size' for more results
        search_results = opensearch_client.search(index="documents", body={
            "size": 50,  #  retrieve more documents
            "query": {
                "knn": {
                    "passage_embedding": {
                        "vector": vector,
                        "k": 50  
                    }
                }
            }
        })
        print(f"Search results: {search_results['hits']['total']['value']} hits")

        
        if search_results['hits']['hits']:
            best_hit = search_results['hits']['hits'][0]  
            best_content = clean_text(best_hit['_source']['content'])
        else:
            best_content = "No relevant document content found."

        print(f"Best content (first 100 chars): {best_content[:100]}...")

        
        combined_input = f"""You are an AI assistant named Coral. Use the following chat history, document content, and the current query to provide a relevant and coherent response that refers to both the chat history and the document.

Chat History:
{chat_history_text}

Document Content:
{best_content}

Current Query: {query_text}

Provide a detailed and relevant answer based on both the user's query and the chat history, along with the document content."""

        
        if "cohere" in model_display_name:
            print("Using a Cohere model")
            max_total_tokens = 4000
            model_type = "cohere"
        elif "llama" in model_display_name or "meta" in model_display_name:
            print("Using a Meta LLaMA model")
            max_total_tokens = 8000
            model_type = "llama"
        else:
            return jsonify({"error": "Unsupported model type"}), 400

        input_tokens = count_tokens(combined_input)
        max_output_tokens = min(max_tokens, max_total_tokens - input_tokens)

        if max_output_tokens <= 0:
            combined_input = truncate_text(combined_input, max_total_tokens - max_tokens)
            input_tokens = count_tokens(combined_input)
            max_output_tokens = max_total_tokens - input_tokens

        print(f"Input tokens: {input_tokens}, Max output tokens: {max_output_tokens}")

        if model_type == "cohere":
            chat_request = oci.generative_ai_inference.models.CohereChatRequest(
                message=combined_input,
                max_tokens=max_output_tokens,
                temperature=temperature,
                frequency_penalty=frequency_penalty,
                top_p=top_p,
                top_k=top_k
            )
        else:  # llama
            content = oci.generative_ai_inference.models.TextContent(text=combined_input)
            message = oci.generative_ai_inference.models.Message(role="USER", content=[content])
            chat_request = oci.generative_ai_inference.models.GenericChatRequest(
                api_format=oci.generative_ai_inference.models.BaseChatRequest.API_FORMAT_GENERIC,
                messages=[message],
                max_tokens=max_output_tokens,
                temperature=temperature,
                frequency_penalty=frequency_penalty,
                top_p=top_p,
                top_k=top_k
            )

        chat_details = oci.generative_ai_inference.models.ChatDetails(
            serving_mode=oci.generative_ai_inference.models.OnDemandServingMode(model_id=model_id),
            chat_request=chat_request,
            compartment_id=app.config['OCI_COMPARTMENT_ID']
        )

        chat_response = generative_ai_inference_client.chat(chat_details)

       
        response_text = handle_ai_response(chat_response, model_type)

        # Store the messages
        store_chat_message(conversation_id, "user", query_text)
        store_chat_message(conversation_id, "ai", response_text.strip())

        return jsonify({"response": response_text.strip(), "conversation_id": conversation_id})

    except Exception as e:
        error_message = f"Error processing query: {str(e)}\n{traceback.format_exc()}"
        app.logger.error(error_message)
        return jsonify({"error": error_message}), 500



def handle_ai_response(chat_response, model_type):
    """
    Extracts and processes the response from the AI model, ensuring the AI leverages retrieved information.
    """
    response_text = ""

    if model_type == "cohere":
        response_text = chat_response.data.chat_response.text
    elif model_type == "llama":
        if chat_response.data.chat_response.choices:
            response_text = chat_response.data.chat_response.choices[0].message.content[0].text

    
    response_text = post_process_response(response_text)

    return response_text


def post_process_response(response_text):
    """
    Removes irrelevant or repetitive content from the AI response and ensures it uses the retrieved data effectively.
    """
    irrelevant_phrases = [
        "Hello, I am Coral, an AI assistant.",
        "How can I help you today?",
        "I am happy to assist you."
    ]

    for phrase in irrelevant_phrases:
        response_text = response_text.replace(phrase, "")

    return response_text.strip()




def test_redis_connection():
    try:
        redis_client.ping()
        print("Successfully connected to Redis")
        mykey_value = redis_client.get('mykey')
        print(f"Value of mykey: {mykey_value}")
        return True
    except redis.exceptions.ConnectionError as e:
        print(f"Failed to connect to Redis: {e}")
        return False
    except Exception as e:
        print(f"An error occurred: {e}")
        return False


if test_redis_connection():
    print("Redis connection test passed")
else:
    print("Redis connection test failed")



if __name__ == '__main__':
    if not opensearch_client.indices.exists(index="documents"):
        create_opensearch_index("documents")
    app.run(host='0.0.0.0', port=5000, debug=True)